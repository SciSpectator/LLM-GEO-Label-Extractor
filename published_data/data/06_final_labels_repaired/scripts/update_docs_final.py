#!/usr/bin/env python3
"""Write the recomputed corpus figures into the manuscript and supplement.

The final corpus is the released 804,427-sample corpus with the re-extracted
samples merged in: same samples, same columns, corrected labels for 59,728 of
them. Every figure written here was produced by the manuscript's OWN evaluation
scripts, run on that merged output, so each replaced number is comparable with
the one it replaces rather than merely newer.

Tables left alone, and why: S1, S3, S11, S13, S19, S20 and main Table 1 report
Phase 1 extraction against fixed benchmarks. Phase 1 is byte-identical to the
released code, and the 1,500-sample benchmarks are unchanged, so their
precision and recall cannot have moved. Their corpus counts shift marginally for
the 5,804 samples Job B re-extracted; that shift is noted in the text rather
than silently written into a benchmark table whose basis this script cannot
reproduce exactly.

Usage:
    update_docs_final.py FIGURES_JSON ACRONYM_JSON SCREEN_JSON \
                         SUPP_IN SUPP_OUT MAIN_IN MAIN_OUT
"""

import json
import sys

import docx


def set_text(par, text):
    if not par.runs:
        par.add_run(text)
        return
    par.runs[0].text = text
    for r in par.runs[1:]:
        r.text = ""


def cells(row, values):
    for c, v in zip(row.cells, values):
        set_text(c.paragraphs[0], str(v))


def find_table(doc, *words):
    for t in doc.tables:
        head = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
        if all(w.lower() in head for w in words):
            return t
    return None


def by_field(rows):
    return {r["field"]: r for r in rows}


def n(x):
    try:
        return f"{int(float(x)):,}"
    except Exception:
        return str(x)


def main() -> int:
    figs_p, acr_p, scr_p, supp_in, supp_out, main_in, main_out = sys.argv[1:8]
    F = json.load(open(figs_p))
    ACR = json.load(open(acr_p))
    SCR = json.load(open(scr_p))

    comp = by_field(F["composition"]["final"])
    voc = by_field(F["vocabulary_precision"]["final"])
    col = by_field(F["collapse"]["final"])
    bran = by_field(F["branch_validity"]["final"])
    cons = by_field(F["consistency"]["final"])
    dbr = F["db_resolution"]["final"]
    unc = F["uncovered"]
    FIELDS = ("Tissue", "Condition", "Treatment")

    # =================== SUPPLEMENT =====================================
    doc = docx.Document(supp_in)
    done = []

    # S2 -- normalization by field
    t = find_table(doc, "raw values", "canonical concepts")
    if t:
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            if f in FIELDS:
                cells(r, [f, n(col[f]["raw_values"]), n(col[f]["canonical"]),
                          n(voc[f]["mesh"]),
                          n(voc[f]["cellosaurus"]) if f == "Tissue" else "n/a",
                          n(voc[f]["oov"])])
        done.append("S2")

    # S5 / S16 -- database resolution
    for t, weighted in ((find_table(doc, "answerable", "precision"), False),
                        (find_table(doc, "weight", "evaluable"), True)):
        if not t:
            continue
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            w = r.cells[1].text.strip().lower() if weighted else "instance"
            key = "instance" if "instance" in w else "value"
            hit = next((x for x in dbr if x["field"] == f and x["weight"] == key), None)
            if not hit:
                continue
            base = [f] + ([r.cells[1].text.strip()] if weighted else [])
            cells(r, base + [n(hit["answerable"]), n(hit["TP"]), n(hit["FP"]),
                             n(hit["FN"]), f"{float(hit['precision']):.3f}",
                             f"{float(hit['recall']):.3f}",
                             f"{float(hit['F1']):.3f}",
                             f"{float(hit['accuracy']):.3f}"])
        done.append("S16" if weighted else "S5")

    # S7 -- per-vocabulary match precision
    t = find_table(doc, "mesh assigned", "mesh verified")
    if t:
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            if f not in FIELDS:
                continue
            v = voc[f]
            mv = int(v["mesh_verified"]) / max(int(v["mesh"]), 1) * 100
            if f == "Tissue":
                cv = int(v["cellosaurus_verified"]) / max(int(v["cellosaurus"]), 1) * 100
                cell_a, cell_v = n(v["cellosaurus"]), f"{n(v['cellosaurus_verified'])} ({cv:.2f}%)"
            else:
                cell_a = cell_v = "n/a"
            io = int(v["oov_inappropriate"]) / max(int(v["oov"]), 1) * 100
            cells(r, [f, n(v["mesh"]), f"{n(v['mesh_verified'])} ({mv:.2f}%)",
                      cell_a, cell_v, n(v["oov"]),
                      f"{n(v['oov_inappropriate'])} ({io:.2f}%)"])
        done.append("S7")

    # S10 -- corpus composition
    t = find_table(doc, "labels present", "not specified")
    if t:
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            if f not in FIELDS:
                continue
            c = comp[f]
            tot = int(c["samples"])
            pres = tot - int(c["not_specified"])
            def pc(x): return f"{n(x)} ({int(x) / tot * 100:.2f}%)"
            cells(r, [f, pc(pres), pc(c["not_specified"]), pc(c["mesh"]),
                      pc(c["cellosaurus"]) if f == "Tissue" else "n/a",
                      pc(c["oov"])])
        done.append("S10")

    # S21 -- what the corpus contains
    t = find_table(doc, "vocabulary", "distinct concepts")
    if t:
        idx = {(r["field"], r["vocabulary"]): r for r in F["resource_summary"]["final"]}

        def same_vocab(short, label):
            """The results file abbreviates what the table spells out."""
            s, l = short.lower(), label.lower()
            if s in l:
                return True
            return s == "oov" and "out-of-vocabulary" in l

        for r in t.rows[1:]:
            f, v = r.cells[0].text.strip(), r.cells[1].text.strip()
            key = next((k for k in idx if k[0] == f and same_vocab(k[1], v)), None)
            if key:
                x = idx[key]
                cells(r, [f, v, n(x["distinct_concepts"]), n(x["samples"]),
                          f"{float(x['pct_of_corpus']):.2f}%"])
        done.append("S21")

    # S22 -- cross-field scope
    t = find_table(doc, "case", "action", "samples")
    if t:
        idx = {r["case"].lower(): r for r in F["field_scope"]["final"]}
        for r in t.rows[1:]:
            key = r.cells[0].text.strip().lower()
            x = idx.get(key)
            if x:
                set_text(r.cells[2].paragraphs[0], n(x["samples"]))
                if x.get("target_field_already_labelled"):
                    set_text(r.cells[3].paragraphs[0],
                             f"{n(x['target_field_already_labelled'])} "
                             f"({float(x['pct_target_labelled']):.2f}%)")
                if len(r.cells) > 4 and x.get("target_field_carries_cellosaurus"):
                    set_text(r.cells[4].paragraphs[0],
                             n(x["target_field_carries_cellosaurus"]))
        done.append("S22")

    # S12 -- acronym behaviour row
    for t in doc.tables:
        for r in t.rows[1:]:
            if r.cells[0].text.strip().lower().startswith("bare disease acronyms"):
                cells(r, ["Bare disease acronyms, resolved per study (repaired)",
                          f"{ACR['fully']} of {ACR['seen']} fully; "
                          f"{ACR['partly']} partly; {ACR['none']} not at all",
                          n(ACR["fn"]),
                          "NSCLC, SLE, GBM, AML resolved; CD ambiguous per study"])
                done.append("S12")
                break

    # S18 -- recall provenance rows
    for t in doc.tables:
        for r in t.rows[1:]:
            lab = r.cells[0].text.strip().lower()
            if lab.startswith("disease-acronym expansio"):
                cells(r, ["Disease-acronym expansion (repaired)",
                          "Curated acronyms, corpus samples",
                          n(ACR["tp"]), n(ACR["fn"]),
                          f"{ACR['tp']}/{ACR['pool']} = {ACR['recall']:.6f}"])
                done.append("S18 acronym row")
            elif lab.startswith("phase 2 ") and "resolu" in lab:
                f = lab.split()[2].capitalize()
                hit = next((x for x in dbr if x["field"] == f
                            and x["weight"] == "instance"), None)
                if hit:
                    tp, fn = int(hit["TP"]), int(hit["FN"])
                    cells(r, [r.cells[0].text.strip(), r.cells[1].text.strip(),
                              n(tp), n(fn),
                              f"{tp}/{tp + fn} = {tp / (tp + fn):.6f}"])
    doc.save(supp_out)
    print("supplement updated:", ", ".join(done))

    # =================== MAIN ===========================================
    m = docx.Document(main_in)
    mdone = []

    t = find_table(m, "labels w/ db answer")
    if t:
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            hit = next((x for x in dbr if x["field"] == f
                        and x["weight"] == "instance"), None)
            if hit:
                cells(r, [f, n(hit["answerable"]), f"{float(hit['precision']):.3f}",
                          f"{float(hit['recall']):.3f}", f"{float(hit['F1']):.3f}",
                          f"{float(hit['accuracy']):.3f}", n(hit["FP"]), n(hit["FN"])])
        mdone.append("Table 2")

    t = find_table(m, "raw values", "canonical")
    if t:
        for r in t.rows[1:]:
            f = r.cells[0].text.strip()
            if f not in FIELDS:
                continue
            cells(r, [f, n(col[f]["raw_values"]), n(col[f]["canonical"]),
                      n(voc[f]["mesh"]),
                      n(voc[f]["cellosaurus"]) if f == "Tissue" else "n/a",
                      n(voc[f]["oov"]), n(unc[f]),
                      f"{float(bran[f]['violation_pct_instances']):.2f}%",
                      f"{float(cons[f]['disagreement_pct']):.2f}%"])
        mdone.append("Table 3")

    for p in m.paragraphs:
        if "collapsed 21,254" in p.text or ("collapsed" in p.text and "canonical concepts" in p.text):
            mesh_tot = sum(int(voc[f]["mesh"]) for f in FIELDS)
            set_text(p, (
                f"Across all 804,427 samples, Phase 2 collapsed "
                f"{n(col['Tissue']['raw_values'])} distinct raw Tissue strings to "
                f"{n(col['Tissue']['canonical'])} canonical concepts, "
                f"{n(col['Condition']['raw_values'])} Condition strings to "
                f"{n(col['Condition']['canonical'])}, and "
                f"{n(col['Treatment']['raw_values'])} Treatment strings to "
                f"{n(col['Treatment']['canonical'])} (Table 3). In total "
                f"{mesh_tot:,} distinct labels resolved to MeSH descriptors and "
                f"{n(voc['Tissue']['cellosaurus'])} Tissue labels to Cellosaurus "
                f"identifiers (e.g. MCF7→CVCL_0031); the remainder were "
                f"assigned to out-of-vocabulary concept clusters. These counts "
                f"are those of the final corpus, in which the samples "
                f"re-extracted after the integrity audit and the acronym "
                f"re-normalisation replace their released labels."))
            mdone.append("collapse paragraph")
            break

    m.save(main_out)
    print("main text updated:", ", ".join(mdone))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
